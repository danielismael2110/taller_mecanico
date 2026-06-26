# -*- coding: utf-8 -*-
"""Genera el Capítulo III - Marco Práctico (GaraGato) en Word."""
import base64
import os
import urllib.request
import zlib

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BRAND = RGBColor(0x5C, 0x0B, 0x8B)
GRIS = RGBColor(0x66, 0x66, 0x66)

doc = Document()

# ---- Estilos base ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

for hname, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)):
    st = doc.styles[hname]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = BRAND
    st.font.bold = True


def h1(t):
    doc.add_heading(t, level=1)


def h2(t):
    doc.add_heading(t, level=2)


def h3(t):
    doc.add_heading(t, level=3)


def parr(texto, justify=True, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(texto)
    r.italic = italic
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(8)
    return p


def vinheta(texto):
    doc.add_paragraph(texto, style="List Bullet")


def numerado(texto):
    doc.add_paragraph(texto, style="List Number")


def tabla(encabezados, filas, anchos=None):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, e in enumerate(encabezados):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(e)
        run.bold = True
        run.font.size = Pt(10)
    for fila in filas:
        celdas = t.add_row().cells
        for i, v in enumerate(fila):
            celdas[i].text = ""
            run = celdas[i].paragraphs[0].add_run(str(v))
            run.font.size = Pt(10)
    if anchos:
        for row in t.rows:
            for i, w in enumerate(anchos):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def kroki_png(puml, nombre):
    """Descarga el PNG de un diagrama PlantUML desde kroki.io. Devuelve la ruta o None."""
    try:
        comp = zlib.compress(puml.encode("utf-8"), 9)
        b64 = base64.urlsafe_b64encode(comp).decode("utf-8")
        url = f"https://kroki.io/plantuml/png/{b64}"
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        ruta = os.path.join(os.path.dirname(__file__), nombre)
        with urllib.request.urlopen(req, timeout=20) as resp, open(ruta, "wb") as f:
            f.write(resp.read())
        return ruta
    except Exception as e:  # noqa
        print("Diagrama no descargado:", e)
        return None


# =====================================================================
# PORTADA DEL CAPÍTULO
# =====================================================================
tit = doc.add_paragraph()
tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tit.add_run("CAPÍTULO III")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = BRAND
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("MARCO PRÁCTICO")
r.bold = True
r.font.size = Pt(18)
parr(
    "En este capítulo se documenta la construcción del sistema web para la gestión y el control del "
    "taller mecánico GaraGato. Se describe la metodología de trabajo adoptada, la planificación del "
    "producto, el modelado de los casos de uso y el desarrollo iterativo organizado en cuatro sprints, "
    "acompañando cada avance con sus pruebas de verificación.",
    italic=True, color=GRIS,
)
doc.add_paragraph()

# =====================================================================
# 3.1 METODOLOGÍA DE DESARROLLO
# =====================================================================
h1("3.1. Metodología de desarrollo")

h2("3.1.1. Enfoque ágil y justificación del uso de Scrum")
parr(
    "El desarrollo de GaraGato se abordó con un enfoque ágil porque el proyecto debía evolucionar a la par "
    "del entendimiento que íbamos ganando sobre la operación real del taller. Desde las primeras "
    "conversaciones quedó claro que los requerimientos no estaban del todo cerrados: la forma en que se "
    "reciben los vehículos, se cotizan los trabajos y se cobran los servicios fue precisándose conforme "
    "mostrábamos avances funcionales. Un método predictivo, que exige congelar el alcance al inicio, habría "
    "chocado con esa realidad."
)
parr(
    "Optamos por Scrum por una razón concreta: su ritmo de entregas cortas permite enseñar resultados "
    "tangibles cada pocas semanas y corregir el rumbo antes de que un malentendido se vuelva costoso. En "
    "lugar de prometer un sistema completo al final, fuimos liberando módulos utilizables —primero la gestión "
    "interna, luego la operación de órdenes y, por último, el portal del cliente—, lo que dio margen para "
    "validar cada pieza con quien conoce el negocio. Scrum también encajó con el tamaño del equipo y con la "
    "necesidad de mantener una lista de prioridades clara y siempre visible."
)

h2("3.1.2. Principios del manifiesto ágil")
parr(
    "El trabajo se guió por los valores del Manifiesto Ágil, que en la práctica se tradujeron en decisiones "
    "cotidianas más que en declaraciones de buenas intenciones:"
)
vinheta("Individuos e interacciones sobre procesos y herramientas: las dudas se resolvían conversando con quien conoce el taller, no llenando formularios.")
vinheta("Software funcionando sobre documentación extensa: al cierre de cada sprint había una versión navegable y desplegable, no solo diagramas.")
vinheta("Colaboración con el cliente sobre negociación contractual: los ajustes de alcance se acordaban mostrando la pantalla real y decidiendo juntos.")
vinheta("Respuesta ante el cambio sobre seguimiento de un plan: cuando surgió, por ejemplo, el flujo de aceptación de órdenes por parte de los mecánicos, se incorporó sin reescribir todo el proyecto.")

h2("3.1.3. Roles en Scrum")
parr(
    "Dado que se trata de un proyecto académico con un equipo reducido, los roles de Scrum se distribuyeron "
    "de forma flexible, pero respetando las responsabilidades que cada uno representa:"
)
tabla(
    ["Rol", "Responsabilidad", "Asignación en el proyecto"],
    [
        ["Product Owner", "Definir y priorizar el Product Backlog; representar la voz del taller.", "Estudiante, con la retroalimentación del propietario del taller."],
        ["Scrum Master", "Facilitar el proceso, retirar impedimentos y cuidar el cumplimiento del marco.", "Estudiante a cargo de la gestión del proyecto."],
        ["Equipo de desarrollo", "Diseñar, construir y probar los incrementos del producto.", "Equipo de desarrollo del proyecto."],
    ],
    anchos=[1.6, 4.0, 2.4],
)

# =====================================================================
# 3.2 PLANIFICACIÓN Y GESTIÓN DEL BACKLOG
# =====================================================================
h1("3.2. Planificación y gestión del backlog")

h2("3.2.1. Creación y estructuración del Product Backlog")
parr(
    "El Product Backlog se construyó a partir de los requerimientos funcionales relevados para el taller. "
    "En lugar de tratarlos como una lista plana, los agrupamos por épicas —grandes bloques de funcionalidad "
    "que comparten un propósito— para poder priorizar de forma ordenada. Cada épica se desglosó luego en "
    "ítems más pequeños, estimables y entregables dentro de un sprint. La prioridad se asignó combinando el "
    "valor para el negocio (qué tan crítico es para operar el taller) y las dependencias técnicas (qué debe "
    "existir antes para que lo demás funcione)."
)

h2("3.2.2. Product Backlog")
parr("El siguiente cuadro resume el Product Backlog del sistema, organizado por épicas y con su prioridad relativa.")
tabla(
    ["ID", "Épica", "Descripción del incremento", "Prioridad"],
    [
        ["E1", "Autenticación y seguridad", "Inicio de sesión, registro de clientes, recuperación de contraseña y control de acceso por rol.", "Alta"],
        ["E2", "Usuarios y roles", "Gestión de usuarios internos: alta, edición, activación y asignación de rol.", "Alta"],
        ["E3", "Clientes y vehículos", "Registro y edición de clientes, sus vehículos e historial de atenciones.", "Alta"],
        ["E4", "Órdenes de trabajo", "Creación, asignación de mecánicos, flujo de estados, diagnóstico, repuestos y adjuntos.", "Alta"],
        ["E5", "Presupuestos", "Generación desde la orden, envío, aprobación o rechazo y versionado.", "Alta"],
        ["E6", "Inventario y compras", "Repuestos, alertas de stock crítico, proveedores y órdenes de compra.", "Media"],
        ["E7", "Citas y agenda", "Agendamiento con horarios disponibles, calendario y bloqueos.", "Media"],
        ["E8", "Pagos y comprobantes", "QR del taller, pago en efectivo, carga y validación de comprobantes.", "Alta"],
        ["E9", "Portal público y del cliente", "Catálogo de servicios, contacto y autogestión del cliente.", "Media"],
        ["E10", "Reportes e indicadores", "Ingresos, servicios más realizados, stock crítico y carga por mecánico.", "Media"],
    ],
    anchos=[0.6, 2.0, 4.3, 1.1],
)

h2("3.2.3. Sprint Backlog")
parr(
    "A partir de la priorización, las épicas se distribuyeron en cuatro sprints. El criterio fue construir "
    "primero los cimientos (datos y seguridad), seguir con la gestión interna y la operación del taller, y "
    "dejar para el final la cara visible al cliente y el cierre de calidad."
)
tabla(
    ["Sprint", "Objetivo", "Épicas / ítems incluidos"],
    [
        ["Sprint 1", "Cimientos del sistema", "Modelo de datos, autenticación base, historias de usuario y mockups."],
        ["Sprint 2", "Gestión interna", "Usuarios y roles; clientes y vehículos; identidad visual."],
        ["Sprint 3", "Operación del taller", "Órdenes de trabajo, presupuestos, inventario, proveedores, compras y servicios."],
        ["Sprint 4", "Cliente y cierre", "Citas, pagos, portal del cliente, reportes, diseño responsive y pruebas."],
    ],
    anchos=[1.1, 2.3, 4.6],
)

# =====================================================================
# 3.3 CASOS DE USO
# =====================================================================
h1("3.3. Casos de uso")

h2("3.3.1. Identificación de actores y requisitos")
parr(
    "Antes de modelar las interacciones se identificaron los actores que intervienen en el sistema. Cada uno "
    "accede a un conjunto de funcionalidades acotado por su rol, lo que más adelante se reforzó a nivel de "
    "base de datos con políticas de seguridad por fila."
)
tabla(
    ["Actor", "Descripción", "Necesidades principales"],
    [
        ["Visitante", "Persona que llega al portal público sin haber iniciado sesión.", "Conocer los servicios, ver la ubicación y registrarse."],
        ["Cliente", "Usuario registrado, dueño de uno o más vehículos.", "Seguir sus órdenes, aprobar presupuestos, pagar y solicitar citas."],
        ["Recepcionista", "Personal de atención y operación diaria del taller.", "Registrar clientes/vehículos, crear órdenes, cotizar, cobrar y agendar."],
        ["Mecánico", "Técnico que ejecuta las reparaciones.", "Aceptar órdenes, registrar diagnóstico y trabajo, usar repuestos."],
        ["Administrador", "Responsable de la configuración y la supervisión.", "Gestionar usuarios, parámetros del taller y consultar reportes."],
    ],
    anchos=[1.5, 3.0, 3.5],
)

h2("3.3.2. Descripción de los casos de uso principales")
parr("Se describen a continuación los casos de uso más representativos del sistema.")
tabla(
    ["Caso de uso", "Actor", "Descripción"],
    [
        ["Registrar cliente y vehículo", "Recepcionista", "Da de alta al cliente y asocia sus vehículos para poder atenderlos."],
        ["Crear orden de trabajo", "Recepcionista", "Registra el ingreso del vehículo, el problema y la cantidad de mecánicos requerida."],
        ["Aceptar orden de trabajo", "Mecánico", "El mecánico acepta o rechaza una orden pendiente de asignación."],
        ["Registrar diagnóstico", "Mecánico", "Documenta hallazgos, trabajo realizado, horas y repuestos utilizados."],
        ["Generar y enviar presupuesto", "Recepcionista", "Crea el presupuesto desde la orden y lo envía al cliente para su respuesta."],
        ["Aprobar o rechazar presupuesto", "Cliente", "El cliente acepta el trabajo o lo rechaza indicando un motivo."],
        ["Subir comprobante de pago", "Cliente", "Adjunta el comprobante de su transferencia por QR para revisión."],
        ["Validar pago", "Recepcionista", "Confirma o anula el comprobante y marca la orden como pagada."],
        ["Solicitar cita", "Cliente", "Reserva un horario disponible para llevar su vehículo."],
        ["Configurar parámetros y reportes", "Administrador", "Ajusta IVA, datos del taller y consulta los indicadores."],
    ],
    anchos=[2.4, 1.6, 4.0],
)

h2("3.3.3. Diagrama de casos de uso")
parr(
    "Los siguientes diagramas, elaborados en notación UML, representan las interacciones entre los actores y "
    "los módulos del sistema, desde la vista general hasta el detalle por área funcional."
)

diagramas = [
    ("Figura 3.1. Vista general de actores y módulos", """@startuml
left to right direction
actor "Visitante" as v
actor "Cliente" as c
actor "Recepcionista" as r
actor "Mecánico" as m
actor "Administrador" as a
package "Sistema Taller GaraGato" {
  usecase "Portal Público" as pp
  usecase "Portal del Cliente" as pc
  usecase "Operación (OT)" as go
  usecase "Recepción y Ventas" as ra
  usecase "Configuración y Reportes" as cr
}
v --> pp
c --> pc
m --> go
r --> ra
a --> cr
@enduml"""),
    ("Figura 3.2. Portal público y portal del cliente", """@startuml
left to right direction
actor "Visitante" as v
actor "Cliente" as c
package "Portal Web" {
  usecase "Ver catálogo de servicios" as uc1
  usecase "Enviar mensaje de contacto" as uc2
  usecase "Registrarse en el sistema" as uc3
  usecase "Solicitar y gestionar citas" as uc4
  usecase "Consultar órdenes de trabajo" as uc5
  usecase "Aprobar/Rechazar presupuestos" as uc6
  usecase "Subir comprobantes de pago" as uc7
}
v --> uc1
v --> uc2
v --> uc3
c --> uc4
c --> uc5
c --> uc6
c --> uc7
@enduml"""),
    ("Figura 3.3. Operaciones: recepcionista y mecánico", """@startuml
left to right direction
actor "Recepcionista" as r
actor "Mecánico" as m
package "Gestión Interna" {
  usecase "Registrar clientes y vehículos" as uc1
  usecase "Crear Órdenes de Trabajo" as uc2
  usecase "Asignar Mecánico a OT" as uc3
  usecase "Registrar diagnóstico y horas" as uc4
  usecase "Generar y enviar presupuestos" as uc5
  usecase "Gestionar inventario y compras" as uc6
  usecase "Validar pagos de clientes" as uc7
}
r --> uc1
r --> uc2
r --> uc3
r --> uc5
r --> uc6
r --> uc7
m --> uc4
m --> uc2 : "Actualiza estado"
m --> uc6 : "Descuenta repuestos"
@enduml"""),
    ("Figura 3.4. Administración", """@startuml
left to right direction
actor "Administrador" as a
package "Panel Administrativo" {
  usecase "Gestionar usuarios internos" as uc1
  usecase "Asignar roles y permisos" as uc2
  usecase "Configurar parámetros (IVA, QR)" as uc3
  usecase "Generar reportes e indicadores" as uc4
}
a --> uc1
a --> uc2
a --> uc3
a --> uc4
@enduml"""),
]

for titulo, puml in diagramas:
    ruta = kroki_png(puml, titulo.split(".")[0].replace(" ", "_") + ".png")
    if ruta:
        doc.add_picture(ruta, width=Inches(5.6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = cap.add_run(titulo)
        rr.italic = True
        rr.font.size = Pt(9)
        rr.font.color.rgb = GRIS
        try:
            os.remove(ruta)
        except OSError:
            pass
    else:
        parr(f"[{titulo} — diagrama UML]", justify=False, italic=True, color=GRIS)

# =====================================================================
# 3.4 SPRINT 1
# =====================================================================
doc.add_page_break()
h1("3.4. Sprint 1: Cimientos del sistema")
parr(
    "El primer sprint se concentró en levantar las bases sobre las que se apoyaría todo lo demás: el modelo "
    "de datos, el esquema de seguridad y los primeros bocetos de interfaz. No produjo pantallas vistosas, "
    "pero sin estos cimientos los sprints siguientes no habrían sido posibles."
)

h2("3.4.1. Historias de usuario")
parr("Las historias de usuario tradujeron las necesidades de cada actor a un lenguaje cercano y verificable.")
tabla(
    ["ID", "Como…", "Quiero…", "Para…"],
    [
        ["HU-01", "Recepcionista", "registrar clientes y sus vehículos", "tener su información lista al crear una orden."],
        ["HU-02", "Recepcionista", "crear una orden de trabajo con el problema reportado", "iniciar formalmente la atención del vehículo."],
        ["HU-03", "Mecánico", "registrar el diagnóstico y el trabajo realizado", "dejar constancia técnica de la reparación."],
        ["HU-04", "Cliente", "aprobar o rechazar un presupuesto", "autorizar el trabajo antes de que se ejecute."],
        ["HU-05", "Cliente", "subir mi comprobante de pago", "que el taller valide mi transferencia."],
        ["HU-06", "Administrador", "configurar el IVA y los datos del taller", "que los cálculos y el portal reflejen la realidad."],
    ],
    anchos=[0.7, 1.7, 3.2, 3.4],
)

h2("3.4.2. Diseño de la base de datos")
parr(
    "El modelo de datos se diseñó sobre PostgreSQL, gestionado a través de Supabase. Se definieron tablas para "
    "perfiles, clientes, vehículos, servicios, repuestos, proveedores, órdenes de compra, órdenes de trabajo y "
    "sus líneas (servicios y repuestos), presupuestos, citas, pagos y notificaciones, además de una tabla de "
    "configuración del taller. Las relaciones se protegieron con reglas de integridad: por ejemplo, un cliente "
    "con órdenes asociadas no puede eliminarse, y un repuesto utilizado en una orden no se borra del catálogo."
)
parr(
    "Buena parte de la lógica de negocio se trasladó a la propia base de datos mediante disparadores y "
    "funciones: la numeración automática de órdenes y comprobantes, el descuento y la devolución de stock al "
    "usar repuestos, el recálculo de totales con IVA y descuento, el bloqueo optimista para evitar ediciones "
    "simultáneas y la generación de notificaciones internas. La seguridad se implementó con políticas de "
    "seguridad a nivel de fila (RLS), de modo que cada rol solo accede a los datos que le corresponden."
)

h2("3.4.3. Elaboración de mockups")
parr(
    "Antes de programar las pantallas definitivas se elaboraron mockups de baja y media fidelidad para acordar "
    "la disposición de la información y el flujo de navegación. Estos bocetos sirvieron para validar, con poco "
    "esfuerzo, decisiones como la estructura del panel con barra lateral colapsable, la ubicación del buscador "
    "y las notificaciones en la cabecera, y la forma de presentar el progreso de una orden mediante un stepper."
)

# =====================================================================
# 3.5 SPRINT 2
# =====================================================================
doc.add_page_break()
h1("3.5. Sprint 2: Gestión interna")
parr(
    "Con los cimientos listos, el segundo sprint entregó los módulos que sostienen la operación administrativa: "
    "la gestión de usuarios internos y la de clientes con sus vehículos. También se fijó la identidad visual "
    "del sistema."
)

h2("3.5.1. Definición de colores")
parr(
    "Se adoptó una paleta basada en un color de marca de tono violeta para los elementos principales, "
    "acompañado de grises neutros para el texto y los fondos, y de colores semánticos para los estados "
    "(verde para éxito, ámbar para advertencia y rojo para error). La paleta se definió mediante variables de "
    "tema, lo que permite mantener la coherencia en toda la aplicación y habilita un modo oscuro sin duplicar "
    "estilos. Cada rol, además, tiene un color distintivo para reforzar la identificación visual."
)

def modulo_pruebas(filas):
    tabla(["Caso de prueba", "Resultado esperado", "Estado"], filas, anchos=[4.0, 3.0, 1.0])

h2("3.5.2. Desarrollo del módulo de usuarios")
parr(
    "El módulo de usuarios permite al administrador gestionar las cuentas internas del taller. Se construyó "
    "siguiendo el patrón que se repetiría en el resto del sistema: un listado con búsqueda, formularios de alta "
    "y edición, y acciones rápidas sobre cada registro."
)
h3("3.5.2.1. Listado de usuarios")
parr("Presenta a los usuarios internos con su nombre, correo, rol y estado, e incluye accesos directos para editar o cambiar su estado.")
h3("3.5.2.2. Creación de usuarios")
parr("Mediante un formulario validado se da de alta a administradores, recepcionistas o mecánicos, asignándoles su rol desde el inicio.")
h3("3.5.2.3. Editar usuario")
parr("Permite actualizar el nombre, el teléfono y el rol de un usuario existente, de modo que sus permisos se ajusten a sus funciones.")
h3("3.5.2.4. Detalle del usuario")
parr("Resume la información de la cuenta y su estado actual, sirviendo de punto de partida para las acciones de gestión.")
h3("3.5.2.5. Pruebas funcionales del módulo de usuarios")
modulo_pruebas([
    ["Crear un usuario con rol recepcionista", "El usuario se registra y aparece en el listado", "Aprobado"],
    ["Editar el rol de un usuario", "El nuevo rol se guarda y se refleja en sus permisos", "Aprobado"],
    ["Desactivar un usuario", "El usuario no puede iniciar sesión", "Aprobado"],
    ["Validar correo duplicado", "El sistema muestra un mensaje claro y no crea la cuenta", "Aprobado"],
])

h2("3.5.3. Desarrollo del módulo de clientes")
parr(
    "El módulo de clientes centraliza la cartera del taller. Desde aquí la recepción registra a los clientes y "
    "consulta su historial, además de exportar el listado cuando se requiere."
)
h3("3.5.3.1. Listado de clientes")
parr("Ofrece una búsqueda con retardo (debounce) por nombre, documento o teléfono, paginación y exportación a Excel.")
h3("3.5.3.2. Nuevo cliente")
parr("Un formulario validado captura los datos de contacto y de identificación del cliente.")
h3("3.5.3.3. Editar cliente")
parr("Permite corregir o completar la información del cliente sin perder su historial.")
h3("3.5.3.4. Detalle del cliente")
parr("Organiza la información en pestañas: datos generales, vehículos asociados e historial de órdenes de trabajo.")
h3("3.5.3.5. Pruebas funcionales del módulo de clientes")
modulo_pruebas([
    ["Registrar un cliente nuevo", "El cliente se guarda y queda disponible para órdenes", "Aprobado"],
    ["Buscar por teléfono", "El resultado aparece en menos de un segundo", "Aprobado"],
    ["Exportar el listado a Excel", "Se descarga un archivo con los clientes", "Aprobado"],
])

h2("3.5.4. Desarrollo del módulo de vehículos")
parr(
    "Cada vehículo se vincula a un cliente y conserva su propio historial de atenciones. El módulo admite "
    "fotografía del vehículo y registra quién dio de alta cada ficha."
)
h3("3.5.4.1. Registro de vehículos")
parr("Captura marca, modelo, año, placa y demás datos, junto con una fotografía opcional.")
h3("3.5.4.2. Edición y detalle del vehículo")
parr("Permite actualizar los datos y muestra el historial de órdenes, alertando si el vehículo tiene saldos pendientes.")
h3("3.5.4.3. Pruebas funcionales del módulo de vehículos")
modulo_pruebas([
    ["Registrar un vehículo con foto", "El vehículo y su imagen se guardan correctamente", "Aprobado"],
    ["Editar un vehículo sin año", "El registro se guarda sin bloquearse por el campo opcional", "Aprobado"],
    ["Ver el historial del vehículo", "Se listan sus órdenes de trabajo", "Aprobado"],
])

# =====================================================================
# 3.6 SPRINT 3
# =====================================================================
doc.add_page_break()
h1("3.6. Sprint 3: Operación del taller")
parr(
    "El tercer sprint dio vida al corazón del sistema: las órdenes de trabajo y todo lo que gira en torno a "
    "ellas. Aquí se integraron la asignación de mecánicos, los presupuestos, el inventario y el catálogo de "
    "servicios."
)

h2("3.6.1. Implementación del sistema de autenticación")
parr(
    "Se consolidó la autenticación apoyada en Supabase Auth, con inicio de sesión por correo y contraseña, "
    "registro público de clientes y recuperación de contraseña. Tras iniciar sesión, el sistema redirige a "
    "cada usuario a su panel según el rol, y cierra la sesión automáticamente tras cinco minutos de "
    "inactividad en el caso de los usuarios internos."
)
h3("3.6.1.1. Panel de inicio")
parr("El panel reúne indicadores y accesos rápidos adaptados al rol, sirviendo como punto de partida del trabajo diario.")
h3("3.6.1.2. Pruebas funcionales de la autenticación")
modulo_pruebas([
    ["Iniciar sesión con credenciales válidas", "Acceso al panel correspondiente al rol", "Aprobado"],
    ["Iniciar sesión con contraseña incorrecta", "Mensaje claro y sin acceso", "Aprobado"],
    ["Inactividad de cinco minutos", "La sesión se cierra automáticamente", "Aprobado"],
])

h2("3.6.2. Desarrollo del módulo de órdenes de trabajo")
parr(
    "La orden de trabajo es el documento central de la operación. La recepción la crea indicando el cliente, el "
    "vehículo, el problema reportado y cuántos mecánicos se requieren; los mecánicos disponibles reciben el "
    "aviso y pueden aceptarla o rechazarla. Cuando se completan los mecánicos requeridos, la orden avanza "
    "automáticamente a diagnóstico. A lo largo de su ciclo de vida, la orden registra cada cambio de estado, "
    "el diagnóstico, los servicios y repuestos —que descuentan stock—, los adjuntos y el cálculo de totales."
)
h3("3.6.2.1. Detalle de la orden y flujo de estados")
parr("El detalle muestra el progreso mediante un stepper, la línea de tiempo de cambios, las tablas de servicios y repuestos, y los mecánicos que trabajan en ella.")
h3("3.6.2.2. Exportación a PDF")
parr("La orden puede exportarse a PDF, incluyendo el problema reportado, el diagnóstico, el trabajo realizado y el desglose de totales.")
h3("3.6.2.3. Pruebas funcionales del módulo de órdenes")
modulo_pruebas([
    ["Crear una orden pendiente de asignación", "Los mecánicos reciben el aviso", "Aprobado"],
    ["Aceptar la orden como mecánico", "La orden avanza a diagnóstico al completar los requeridos", "Aprobado"],
    ["Agregar un repuesto a la orden", "El stock se descuenta automáticamente", "Aprobado"],
    ["Exportar la orden a PDF", "Se descarga el documento con el detalle completo", "Aprobado"],
])

h2("3.6.3. Desarrollo del módulo de presupuestos")
parr(
    "Los presupuestos se generan directamente desde la orden, tomando sus servicios y repuestos como líneas. "
    "Una vez enviado, el cliente lo revisa en su portal y lo aprueba o lo rechaza con un motivo obligatorio; "
    "en ambos casos el taller recibe la notificación y el estado de la orden se actualiza. El sistema conserva "
    "el versionado de los presupuestos."
)
h3("3.6.3.1. Pruebas funcionales del módulo de presupuestos")
modulo_pruebas([
    ["Generar presupuesto desde una orden", "Se crean las líneas con sus totales", "Aprobado"],
    ["Enviar el presupuesto al cliente", "El cliente recibe la notificación", "Aprobado"],
    ["Rechazar sin motivo", "El sistema exige el motivo y no continúa", "Aprobado"],
])

h2("3.6.4. Desarrollo del módulo de inventario, proveedores y compras")
parr(
    "El inventario controla los repuestos del taller con sus precios, ubicación y stock mínimo. El sistema "
    "señala los repuestos en stock crítico y notifica al personal. Los movimientos quedan registrados, y los "
    "ajustes manuales exigen un motivo. El módulo se complementa con la gestión de proveedores y con órdenes "
    "de compra cuyo estado se administra (borrador, enviada, recibida); al recibir los ítems, el stock se "
    "incrementa automáticamente."
)
h3("3.6.4.1. Pruebas funcionales del módulo de inventario")
modulo_pruebas([
    ["Registrar un repuesto", "El repuesto queda disponible para las órdenes", "Aprobado"],
    ["Ajustar stock sin motivo", "El sistema exige el motivo del ajuste", "Aprobado"],
    ["Recibir una orden de compra", "El stock aumenta según lo recibido", "Aprobado"],
])

h2("3.6.5. Desarrollo del módulo de servicios")
parr(
    "El catálogo de servicios alimenta tanto el portal público como las órdenes de trabajo. La recepción puede "
    "crear, editar y eliminar servicios, asignarles imagen y precio, y decidir cuáles son visibles en el portal "
    "y cuáles se destacan en la página de inicio."
)
h3("3.6.5.1. Pruebas funcionales del módulo de servicios")
modulo_pruebas([
    ["Crear un servicio visible en el portal", "Aparece en el catálogo público", "Aprobado"],
    ["Marcar un servicio como destacado", "Se muestra en la página de inicio", "Aprobado"],
])

# =====================================================================
# 3.7 SPRINT 4
# =====================================================================
doc.add_page_break()
h1("3.7. Sprint 4: Cliente, pagos y cierre")
parr(
    "El último sprint completó la cara visible para el cliente y cerró el proyecto con un trabajo de calidad: "
    "citas, pagos, portal del cliente, reportes, diseño responsive y pruebas."
)

h2("3.7.1. Citas y agenda")
parr(
    "El módulo de citas muestra únicamente los horarios disponibles, en intervalos de media hora, marcando en "
    "gris los ocupados o fuera de horario para que el cliente entienda por qué no puede elegirlos. La base de "
    "datos impide que un mecánico tenga dos citas a la vez, y el administrador puede bloquear días u horarios "
    "por feriados o mantenimiento. El personal visualiza la agenda en un calendario semanal con colores por "
    "estado."
)

h2("3.7.2. Pagos y comprobantes")
parr(
    "El cobro contempla dos medios: efectivo y QR. El cliente ve sus órdenes con saldo pendiente, escanea el "
    "QR del taller y sube su comprobante, que queda en revisión hasta que la recepción lo valida o lo anula. "
    "Al cubrirse el total, la orden se marca como pagada. Cada cambio relevante genera una notificación interna."
)

h2("3.7.3. Portal del cliente")
parr(
    "El portal del cliente reúne en un solo lugar su actividad: sus vehículos, la orden activa con su progreso, "
    "la próxima cita y accesos para aprobar presupuestos, pagar y solicitar citas. La intención fue que el "
    "cliente pueda seguir su reparación sin llamar al taller."
)

h2("3.7.4. Reportes e indicadores")
parr(
    "Para la toma de decisiones, el administrador dispone de reportes de ingresos por día con filtros de fecha, "
    "el top de servicios más realizados, los repuestos en stock crítico y la carga de trabajo por mecánico. "
    "Todos pueden exportarse a PDF y Excel."
)

h2("3.7.5. Diseño responsive")
parr(
    "La interfaz se construyó con un enfoque adaptable. En pantallas grandes el panel muestra la barra lateral "
    "fija; en dispositivos móviles, esa barra se oculta tras un menú accesible con un botón, y las tablas y "
    "formularios reorganizan sus columnas para seguir siendo usables en anchos reducidos."
)

h2("3.7.6. Pruebas")
h3("3.7.6.1. Pruebas funcionales de extremo a extremo")
parr(
    "Se ejecutaron pruebas de extremo a extremo sobre los flujos críticos —inicio de sesión, creación y "
    "atención de una orden, aprobación de presupuesto y pago— recorriendo la aplicación tal como lo haría un "
    "usuario real, para confirmar que los módulos funcionan de forma integrada y no solo de manera aislada."
)
h3("3.7.6.2. Evaluación de la velocidad de carga")
parr(
    "Se midió el rendimiento de las pantallas principales verificando que la carga se mantuviera dentro de "
    "tiempos aceptables. La generación de documentos PDF se resolvió del lado del cliente para no recargar el "
    "servidor, y las búsquedas se optimizaron con un retardo que evita consultas innecesarias."
)
h3("3.7.6.3. Análisis e interpretación de métricas")
parr(
    "Los resultados se interpretaron a la luz de los requerimientos no funcionales: tiempos de respuesta, "
    "claridad de los mensajes de error, comportamiento en distintos tamaños de pantalla y exactitud de los "
    "cálculos monetarios. Los hallazgos guiaron pequeños ajustes de usabilidad y desempeño."
)
h3("3.7.6.4. Conclusión de las pruebas")
parr(
    "Las pruebas confirmaron que el sistema cumple con los flujos previstos y con los criterios de calidad "
    "definidos. Las observaciones encontradas durante esta etapa se corrigieron antes del cierre, dejando una "
    "versión estable y lista para su uso en el taller."
)

# ---- Pie de numeración ----
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("Capítulo III - Marco Práctico · GaraGato")
run.font.size = Pt(8)
run.font.color.rgb = GRIS

ruta_salida = os.path.join(os.path.dirname(__file__), "Capitulo_III_Marco_Practico.docx")
doc.save(ruta_salida)
print("OK ->", ruta_salida)
